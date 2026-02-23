import requests
import os
import pandas as pd
import json
from datetime import datetime
import sys
import subprocess
import re
from datetime import datetime
from openai import OpenAI
from zyte_api import ZyteAPI
from docx import Document

#-----------------------Prerequisites Function----------------------
class NYTApiError(Exception):
    """Custom exception for NYT API errors.
    Provides friendly error messages for common issues
    like invalid keys, network failures, and rate limits."""
    pass
# Try to import zyte_api, auto-install if missing. zyte_api is used to extract the content of the articles.
try:
    from zyte_api import ZyteAPI
except ImportError:
    print("=" * 60)
    print("zyte_api not found. Attempting to install...")
    print(f"Using Python: {sys.executable}")
    print("=" * 60)
    
    # Try to install using the current Python interpreter
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "zyte-api"])
        print("✅ Installation successful! Re-importing...")
        from zyte_api import ZyteAPI
        print("✅ zyte_api imported successfully!")
    except subprocess.CalledProcessError:
        print("❌ Auto-installation failed.")
        print()
        print("Please run this command manually:")
        print(f"  {sys.executable} -m pip install zyte-api")
        print()
        print("Or configure your IDE to use this Python:")
        print("  C:\\Users\\user\\AppData\\Local\\Programs\\Python\\Python314\\python.exe")
        sys.exit(1)
# Try to import docx, auto-install if missing
try:
    from docx import Document
except ImportError:
    import sys
    import subprocess
    print("=" * 60)
    print("ERROR: docx module not found!")
    print("=" * 60)
    print(f"Current Python: {sys.executable}")
    print()
    print("Attempting to install python-docx...")
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        print("✅ Installation successful! Re-importing...")
        from docx import Document
        print("✅ docx imported successfully!")
    except subprocess.CalledProcessError:
        print("❌ Auto-installation failed.")
        print()
        print("Please run this command manually:")
        print(f"  {sys.executable} -m pip install python-docx")
        sys.exit(1)

def load_env_file(filepath=".env"):
    with open(filepath, "r") as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                key, value = line.split("=", 1)
                os.environ[key.strip()] = value.strip()
    return True
# Function to remove image descriptions from the article content
def remove_image_descriptions(text: str) -> str:
    """
    Remove image descriptions and captions from article text.
    Filters out common patterns like "This image shows...", "Photo by...", etc.
    
    Parameters:
        text: The article text to clean
    
    Returns:
        str: Cleaned text without image descriptions
    """
    
    if not text:
        return ""
    
    lines = text.split('\n')
    cleaned_lines = []
    
    # Patterns that indicate image descriptions
    image_patterns = [
        r'^This image',
        r'^This photo',
        r'^Photo by',
        r'^Image by',
        r'^Image released by',
        r'^Image via',
        r'^Photo via',
        r'^\(Photo by',
        r'^\(Image by',
        r'^\(AP Photo',
        r'^\(.*?via AP\)',  # (Something via AP)
        r'^\(.*?Photo.*?\)',  # (Something Photo Something)
        r'^Director.*?pose for photographers',  # Photo call descriptions
        r'^This image released by',
    ]
    
    for line in lines:
        line = line.strip()
        if not line:
            cleaned_lines.append('')
            continue
        
        # Skip lines that match image description patterns
        is_image_desc = False
        for pattern in image_patterns:
            if re.match(pattern, line, re.IGNORECASE):
                is_image_desc = True
                break
        
        # Skip very short lines that are likely captions (less than 20 chars)
        # unless they look like actual content
        if not is_image_desc and len(line) < 20:
            # Keep if it looks like a sentence (ends with punctuation)
            if not re.search(r'[.!?]$', line):
                # Skip if it's all caps (likely a caption)
                if line.isupper():
                    continue
        
        if not is_image_desc:
            cleaned_lines.append(line)
    
    # Join lines and clean up multiple newlines
    cleaned_text = '\n'.join(cleaned_lines)
    # Remove excessive blank lines (more than 2 consecutive)
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
    
    return cleaned_text.strip()
# Get full article data -> parse article data
def parse_article(article: dict) -> dict:
    """Parse a single article from the NYT API response.
    Extracts title, date, section, url"""
    return {
        "title": article.get("title", "N/A"),
        "published_date": article.get("published_date", "N/A"),
        "section": article.get("section", "N/A"),
        "url": article.get("url", "N/A"),
    }
# Valid endpoint types and period options
VALID_ENDPOINTS = {"viewed": "Most Viewed", "emailed": "Most Emailed", "shared": "Most Shared"}
VALID_PERIODS = [1, 7, 30]
def extract_articles(endpoint: str = "viewed", period: int =1, num_articles: int =20):
    """Fetch most popular articles from the NYT API.
    
    Parameters:
        endpoint: One of 'viewed', 'emailed', 'shared'
        period: Time period in days (1, 7, or 30)
        num_articles: Number of articles to return (1-20)
        api_key: NYT API key (if None, loads from .env)
    
    Returns:
        List of parsed article dictionaries
    
    Raises:
        NYTApiError: On any API or network error with a friendly message
    """
    if endpoint not in VALID_ENDPOINTS:
        raise ValueError("Invalid endpoint. Please choose from 'viewed', 'emailed', or 'shared'.")
    if period not in VALID_PERIODS:
        raise ValueError("Invalid period. Please choose from 1, 7, or 30 days.")

    # Get API key
    load_env_file()
    api_key = os.getenv("TEST_API_KEY")
    if not api_key:
        raise ValueError("TEST_API_KEY not found in .env file. Please set it up first.")
    url = f"https://api.nytimes.com/svc/mostpopular/v2/{endpoint}/{period}.json"
    params = {"api-key": api_key}
    # Make the request with error handling
    try:
        response = requests.get(url, params=params, timeout=15)
    except requests.ConnectionError:
        raise NYTApiError("Network error: Could not connect to the NYT API. Check your internet connection.")
    except requests.Timeout:
        raise NYTApiError("Request timed out. The NYT API is not responding. Please try again.")
    except requests.RequestException as e:
        raise NYTApiError(f"Request failed: {str(e)}")
    # Handle HTTP error codes with friendly messages
    if response.status_code == 401:
        raise NYTApiError("Invalid API key. Please check your TEST_API_KEY in the .env file.")
    elif response.status_code == 403:
        raise NYTApiError("Access forbidden. Your API key may not have access to this endpoint.")
    elif response.status_code == 429:
        raise NYTApiError("Rate limit exceeded. Please wait a moment and try again.")
    elif response.status_code != 200:
        raise NYTApiError(f"API returned an error (HTTP {response.status_code}). Please try again later.")

    # Parse JSON response
    try:
        data = response.json()
    except (json.JSONDecodeError, ValueError):
        raise NYTApiError("Could not parse the API response.")

    # Parse each article and limit to requested number
    articles = [parse_article(article) for article in data["results"][:num_articles]] 
    return articles
# Extract the content of the article using Zyte API
def extract_article_content(urls: list[str]):
    # Use Zyte API to extract the content of the article
    client = ZyteAPI(api_key=os.getenv("ZYTE_API_KEY"))
    contents = []
    for url in urls:
        try:
            content = client.get({
                "url": url,
                "article": True
            })
            # Extract the only the content of the article
            content = content.get("article", {}).get("articleBody", "")
            content = remove_image_descriptions(content)
            contents.append(content)
        except Exception as e:
            print(f"Error extracting content from {url}: {e}")
            return None
    return contents
# The function to summarize the content of the article using customized Ollama model
#Schema for the response from the Ollama model - structured output with key insights only
schema = {
    "type": "object",
    "properties": {
        "key_insights": {
            "type": "array",
            "description": "List of key insights and important points extracted from the article. Each insight should be a concise, actionable point.",
            "items": {
                "type": "string"
            },
            "minItems": 3,
            "maxItems": 7
        }
    },
    "required": ["key_insights"]
}
def ollama_summarize(contents: list[str]):
    """
    Input: list of article contents
    Output: list of dictionaries with key insights
    """
    PORT = 11434
    OLLAMA_HOST = f"http://localhost:{PORT}"
    url = f"{OLLAMA_HOST}/api/generate"
    insights = []
    #Build the request body as a dictionary
    for content in contents:
        body = {
            "model": "newsanalyst:latest",
            "prompt": content,
            "format": schema,
            "stream": False
        }
        #Build and send the POST request to the Ollama REST API
        response = requests.post(url, json=body)
        response = response.json()
        insight = response.get("response", "").strip()
        if isinstance(insight, str) and insight:
            try:
                parsed = json.loads(insight)
            except json.JSONDecodeError:
                parsed = {"key_insights": [insight] if insight else []}
        else:
            parsed = insight if isinstance(insight, dict) else {"key_insights": []}
        insights.append(parsed.get("key_insights", []))
    return insights
# Generate a report of the article using OPENAI API
def generate_report(data: pd.DataFrame):
    """
    Input: dataframe with title, published_date, section, url, and insights
    Output: report of the article
    """
    articles = json.dumps(data.to_dict(orient="records"), indent=2, ensure_ascii=False)
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    input_messages = [
        {
            "role": "user",
            "content": articles
        }
    ]
    response = client.responses.create(
        model="gpt-5",
        instructions = """
        You are a senior news analyst with 30 years of experience in editorial trend reporting.

        You will receive multiple articles as a JSON array. Each item contains:
        - title
        - published_date
        - section
        - url
        - key_insights

        Task:
        Synthesize the items into ONE consolidated news trend brief of 300–400 words.

        Analytical requirements:
        - Identify overarching themes and emerging trends across the set.
        - Highlight shifts in policy, market direction, and public sentiment where supported by the inputs.
        - Do NOT summarize articles one by one; write as a synthesis.
        - Do NOT reference or include URLs.
        - Do NOT mention “the provided articles” or “the input data”.
        - Maintain an objective, analytical, professional tone. Avoid hype and speculation.

        Date handling:
        - Determine the reporting date range from published_date values.
        - Use the earliest and most recent dates (inclusive). If only one date exists, use that single date.

        Output format (exact headings and order; no extra sections):
        # Data Report
        ### Date: <earliest published_date> to <most recent published_date>
        ### Summary:
        <300–400 word consolidated trend brief>
        ### Sources included:
        - <title 1> <url 1>
        - <title 2> <url 2>
        - <title 3> <url 3>
        """,
        input=input_messages,  # Pass as array of message objects with roles
        stream=False
    )
    return response.output_text
# Save the report to a file
def save_report(report: str):
    #save the data_report to a file
    doc = Document()
    for line in report.split("\n"):
        if line.startswith("# "):
            # Main heading
            doc.add_heading(line[2:], level=1)
        elif line.startswith("### "):
            # Subheading
            doc.add_heading(line[3:], level=3)
        elif line.startswith("- "):
            # Bullet point
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            # Regular paragraph
            doc.add_paragraph(line)
    doc.save("data_report.docx")
    print("✅Data report saved to data_report.docx")

def main():
    data = extract_articles(endpoint="viewed", period=7, num_articles=30)
    content = extract_article_content([article["url"] for article in data])
    insights = ollama_summarize(content)
    # Combine the data and insights into a dataframe
    df = pd.DataFrame(data)
    df["insights"] = insights
    # Generate a report of the article using OPENAI API
    report = generate_report(df)
    save_report(report)

if __name__ == "__main__":
    main()